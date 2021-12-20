from docx import Document
from docx.shared import Pt, RGBColor, Parented
from docx.oxml.ns import qn, nsdecls
import cv2
import numpy as np
import os
from docx.oxml import parse_xml
import pandas as pd

PATH = "C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령증_2021_전측방시트.xlsx"
document = Document()
data = pd.read_excel(PATH)
cnt = 1
for exp_name in data['시험명']:
    for (path, dirs, files) in os.walk('C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령이미지_파일명검수버전'):
        if path.split('\\')[-1] == exp_name:
            print(cnt)
            cnt = cnt + 1
            para = document.add_paragraph()
            run = para.add_run('    전측방 학습용 영상 Annotation 작업 수당 수령증 별첨 증빙자료')
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(14)
            run.bold = True

            para = document.add_paragraph()
            run = para.add_run('시험명 : {}'.format(exp_name))

            para = document.add_paragraph()
            run = para.add_run()
            for file in os.listdir(path + '\\JungBing'):
                img_array = np.fromfile(path + '\\JungBing\\' + file, np.uint8)
                img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)
                dst = cv2.resize(img, dsize=(320, 180), interpolation=cv2.INTER_AREA)
                cv2.imwrite("C:\\Users\\jcy37\\Desktop\\dst.jpg", dst)
                run.add_picture("C:\\Users\\jcy37\\Desktop\\dst.jpg")
            month = str(data[data['시험명'] == exp_name]['월'].values[0])
            day = str(int(data[data['시험명'] == exp_name]['일'].values[0]))
            if len(day) == 1:
                day = '0' + day
            document.save('C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\별첨\\' + data[data['시험명'] == exp_name]['작업자'].values[0] + '_' +
                          month + day + ".docx")
            document = Document()
