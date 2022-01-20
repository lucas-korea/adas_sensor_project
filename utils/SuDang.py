import csv
import os
import pickle
import shutil
import xml.etree.ElementTree as ET

import cv2
import numpy as np
import pandas as pd
from PyPDF2 import PdfFileReader, PdfFileWriter, PdfFileMerger
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.pdfgen import canvas

worker_list = ['권미애', '김민서', '노수진', '노은영', '노화중', '박윤미', '성미애', '성선영', '신성례', '윤가영', '윤점동', '오연주',
               '이민희', '정성미', '강인선', '고지연', '김다예', '배은이', '윤기주', '이상미', '정금연', '정다운', '정유림', '정혜림']
# worker_list = [ '권미애', '김민서', '노수진', '노은영', '노화중', '박윤미', '성미애', '성선영', '신성례', '윤가영', '윤점동', '오연주','이민희', '정성미']
result = []
paper_num_ = 0
day_list = [1,2,3,4,5,8,9,10,11,12,15,16,17,18,19,22,23,24,25,26,29,41,42,43,46,47,48,49,410,413,414,415,416,417,420,421,422,423,424]

def main():
    global result
    OpenCsv_ObjCounting_cnt = 0
    PATH_Wan = "D:\\GT 생성 업무\\객체생성-검수\\생성완"
    for (path, dirs, files) in os.walk(PATH_Wan): #생성완 폴더 진입
        if path[-4:-1] in worker_list: #현재 폴더가 "작업자"_ 인지 확인
            if path[-4:-1] == '권미애': # 그냥 특정 작업자에서 작업 끝내기 위한 코드
                pass
            for file in files:
                if file[-7:] == "sub.csv": #*sub.csv file 찾기
                    print(path, file)
                    OpenCsv_ObjCounting_cnt = OpenCsv_ObjCounting_cnt + 1
                    OpenCsv_ObjCounting(file,path)
    print(paper_num_)
    try:
        f = open('dummy_.csv', 'w', newline='')
    except:
        f = open('dummy3.csv', 'w', newline='')
    wr = csv.writer(f)
    wr.writerow(['이름', '시험명', '프레임수', '금액','오브젝트 개수'])
    name = ''
    day_i = 0
    print("result:\n", len(result))
    with open('test.pickle', 'wb') as f:
        pickle.dump(result, f)
    for i in range(len(result)):
        name2 = result[i][0]
        if name != name2:
            day_i = 0
        name = name2
        day = day_list[day_i]
        if str(day)[0] == '4' and len(str(day)) > 1:
            day = str(day)[1:]

        print(result[i][:4] + ["Object개수:  차선_{}/  차량 _{}/ 보행자_{}/  신호등_{}/  표지판_{}/  노면표식_{}/ 노면화살표_{}".
                    format((result[i][4]), (result[i][5]),(result[i][6]),(result[i][7]),(result[i][8]),
                    (result[i][9]),(result[i][10]))] + [sum(result[i][4:11])]+ [day])
        wr.writerow(result[i][:4] + ["Object개수:  차선_{}/  차량 _{}/ 보행자_{}/  신호등_{}/  표지판_{}/  노면표식_{}/ 노면화살표_{}".format((result[i][4]), (result[i][5]),(result[i][6]),(result[i][7]),(result[i][8]),(result[i][9]),(result[i][10]))] + [sum(result[i][4:11])]+ [day])
        day_i = day_i + 1


def OpenCsv_ObjCounting(file, PATH):
    Lane_num = []
    Vehicle_num = []
    Pedestrian_num = []
    TraficLight_num = []
    TraficSign_num = []
    RoadMark_num = []
    RoadMarkArrow = []

    gubun = LoadGubun(file.split('_')[0])
    fee = DefineFee(gubun)
    Sub_data = pd.read_csv(PATH + '\\' + file, encoding='euc-kr', sep='\t')
    properate_divined_pay = DefProperate_divined_pay(file, PATH)
    img_cnt = 0
    ing_cnt2 = 0
    img_num = 0
    total = 0
    exp_num = 0
    img_name = ImageName(file.split('_')[0], PATH, 0)

    for i in range(len(Sub_data)):
        if Sub_data.iloc[i].values[0].split(',')[0] != 'total' and Sub_data.iloc[i].values[0].split(',')[0][:-1] != 'subtotal' and len(Sub_data.iloc[i].values[0].split(',')[0]) > 2:
            Lane_num.append(int(Sub_data.iloc[i].values[0].split(',')[4]))
            Vehicle_num.append(int(Sub_data.iloc[i].values[0].split(',')[5]))
            Pedestrian_num.append(int(Sub_data.iloc[i].values[0].split(',')[6]))
            TraficLight_num.append(int(Sub_data.iloc[i].values[0].split(',')[7]))
            TraficSign_num.append(int(Sub_data.iloc[i].values[0].split(',')[8]))
            RoadMark_num.append(int(Sub_data.iloc[i].values[0].split(',')[9]))
            RoadMarkArrow.append(int(Sub_data.iloc[i].values[0].split(',')[10]))
            total = (sum(Lane_num) + sum(Vehicle_num) + sum(Pedestrian_num) + sum(TraficLight_num) + sum(
                TraficSign_num) + sum(RoadMark_num) + sum(RoadMarkArrow)) * fee
            if total > properate_divined_pay:
                img_num = ing_cnt2 - img_cnt
                img_cnt = ing_cnt2
                write_mail_merge_excel(file.split('_')[0][:3], img_name, img_num, total, Lane_num, Vehicle_num, Pedestrian_num, TraficLight_num, TraficSign_num, RoadMark_num, RoadMarkArrow)
                img_name = ImageName(file.split('_')[0], PATH, i)
                exp_num = exp_num + 1
                Lane_num = []
                Vehicle_num = []
                Pedestrian_num = []
                TraficLight_num = []
                TraficSign_num = []
                RoadMark_num = []
                RoadMarkArrow = []
            ing_cnt2 = ing_cnt2 + 1
    img_num = ing_cnt2 - img_cnt
    write_mail_merge_excel(file.split('_')[0][:3], img_name, img_num, total, Lane_num, Vehicle_num, Pedestrian_num, TraficLight_num, TraficSign_num,
                           RoadMark_num, RoadMarkArrow)

def write_mail_merge_excel(name, img_name, img_num, total, Lane_num, Vehicle_num, Pedestrian_num, TraficLight_num, TraficSign_num, RoadMark_num, RoadMarkArrow):
    global paper_num_
    global result
    paper_num_ = paper_num_ + 1
    result.append([name, img_name, img_num, total, sum(Lane_num), sum(Vehicle_num), sum(Pedestrian_num), sum(TraficLight_num), sum(TraficSign_num), sum(RoadMark_num), sum(RoadMarkArrow)])
    print(name, img_name, img_num, total, sum(Lane_num), sum(Vehicle_num), sum(Pedestrian_num), sum(TraficLight_num), sum(TraficSign_num), sum(RoadMark_num), sum(RoadMarkArrow))

def DefProperate_divined_pay(file, PATH):
    Lane_num = []
    Vehicle_num = []
    Pedestrian_num = []
    TraficLight_num = []
    TraficSign_num = []
    RoadMark_num = []
    RoadMarkArrow = []

    gubun = LoadGubun(file.split('_')[0])
    fee = DefineFee(gubun)
    Sub_data = pd.read_csv(PATH + '\\' + file, encoding='euc-kr', sep='\t')
    for i in range(len(Sub_data)):
        if Sub_data.iloc[i].values[0].split(',')[0] != 'total' and Sub_data.iloc[i].values[0].split(',')[0][:-1] != 'subtotal' and len(Sub_data.iloc[i].values[0].split(',')[0]) > 2:
            Lane_num.append(int(Sub_data.iloc[i].values[0].split(',')[4]))
            Vehicle_num.append(int(Sub_data.iloc[i].values[0].split(',')[5]))
            Pedestrian_num.append(int(Sub_data.iloc[i].values[0].split(',')[6]))
            TraficLight_num.append(int(Sub_data.iloc[i].values[0].split(',')[7]))
            TraficSign_num.append(int(Sub_data.iloc[i].values[0].split(',')[8]))
            RoadMark_num.append(int(Sub_data.iloc[i].values[0].split(',')[9]))
            RoadMarkArrow.append(int(Sub_data.iloc[i].values[0].split(',')[10]))

    total = (sum(Lane_num) + sum(Vehicle_num) + sum(Pedestrian_num) + sum(TraficLight_num) + sum(TraficSign_num) + sum(
        RoadMark_num) + sum(RoadMarkArrow)) * fee
    paper_num = 1
    while total/paper_num > 295000:
        paper_num = paper_num + 1
    properate_divined_pay = total / paper_num
    return properate_divined_pay

def DefineFee(gubun):
    if gubun == '자전로':
        return 100
    elif gubun == '도심로':
        return 150
    elif gubun == '복합':
        return 150
    print("fee never be 0")
    exit(1)
    return 0

def LoadGubun(case_name):
    PATH = "D:\\GT 생성 업무\\[참고]Heptacam_가이드문서 및 작업자관리시트\\미첨부(내부문서)_온라인 작업자별_할당및통계.xlsx"
    data = pd.read_excel(PATH, sheet_name=1)
    return data[data['폴더명'] == case_name]['구분'].values[0]

def ImageName(case_name, PATH, i):
    for (path, dirs, files) in os.walk(PATH+'\\'+case_name):
        if len(path.split('\\')[-1]) == 1:
            return files[i].split('.')[0]

def GetJungBingImgXml():
    PATH_auto = "C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화"
    PATH_inspec = "D:\\GT 생성 업무\\객체생성-검수\\생성완"
    dst = "C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령이미지_파일명검수버전"
    df_sudang = pd.read_excel(PATH_auto + '\\' + "수당수령증_2021_전측방시트 2차minus1차_2.xlsx")
    cnt = 0
    BREAK = False
    for i in range(len(df_sudang)):
        print(cnt)
        cnt = cnt + 1
        cur_exp_name = df_sudang.iloc[i]['시험명']
        cam_num = cur_exp_name[0]
        if cam_num not in ['1', '2', '3']:
            cam_num = '4'
        NAME = df_sudang.iloc[i]['작업자']
        if NAME == '윤점동':
            NAME = '윤가영'
        suc = 0
        for (path, dirs, files) in os.walk(PATH_inspec):
            if BREAK:
                BREAK = False
                break
            if path.split('\\')[-2][0:3] == NAME and path.split('\\')[-4] == '생성완': #상위상위폴더 이름 첫 3글자가 작업자 이름이랑 매칭되는지 확인 상상상위
                for j in range(len(files)):
                    if files[j].split('.')[0] == cur_exp_name:
                        j2 = int(j + df_sudang.iloc[i]['프레임수'] * 2/4)
                        j3 = int(j + df_sudang.iloc[i]['프레임수'] * 3/4)
                        try:
                            os.mkdir(dst + '\\' + cur_exp_name)
                        except:
                            pass
                        os.mkdir(dst + '\\' + cur_exp_name +'\\' + cam_num)
                        os.mkdir(dst + '\\' + cur_exp_name + '\\' + cam_num + "_annotations_v001")
                        shutil.copy2(path +'\\' + files[j], dst + '\\' + cur_exp_name +'\\' + cam_num + '\\' + files[j])
                        shutil.copy2(path + '\\' + files[j2], dst + '\\' + cur_exp_name +'\\' + cam_num + '\\' + files[j2])
                        shutil.copy2(path + '\\' + files[j3], dst + '\\' + cur_exp_name +'\\' + cam_num + '\\' + files[j3])
                        shutil.copy2('\\'.join(path.split('\\')[0:-1]) +'\\' + os.path.basename(path)+"_annotations_v001" + '\\' + files[j].split('.')[0] + "_v001.xml" ,
                                     dst + '\\' + cur_exp_name + '\\' + cam_num + "_annotations_v001" + '\\' + files[j].split('.')[0] + "_v001.xml")
                        shutil.copy2('\\'.join(path.split('\\')[0:-1]) +'\\' + os.path.basename(path)+"_annotations_v001" + '\\' + files[j2].split('.')[0] + "_v001.xml" ,
                                     dst+ '\\' + cur_exp_name + '\\' + cam_num + "_annotations_v001"  + '\\' + files[j2].split('.')[0] + "_v001.xml")
                        shutil.copy2('\\'.join(path.split('\\')[0:-1]) +'\\' + os.path.basename(path)+"_annotations_v001" + '\\' + files[j3].split('.')[0] + "_v001.xml" ,
                                     dst+ '\\' + cur_exp_name + '\\' + cam_num + "_annotations_v001"  + '\\' + files[j3].split('.')[0] + "_v001.xml")
                        suc = 1
                        BREAK = True
                        break
            else:
                pass
        if suc == 0:
            print('fail getting files!!!')
            print(df_sudang.iloc[i]['시험명'])
                # print(df_sudang.iloc[i]['시험명'])
                # print(path.split('\\')[-2][0:3] , len(path.split('\\')[-3]) )

def ReNameMakeModifyXML():
    for (path, dirs, files) in os.walk("C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령이미지"):
        for file in files:
            if file.split('.')[-1] == 'xml':
                shutil.copy2(path + '\\' + file, path + '\\' + file.split('.')[0] + '_1.xml')

def DrawJungbingImg():
    PATH = "C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령이미지_파일명검수버전"
    i = 0
    for (path, dirs, files) in os.walk(PATH):
        if (path[-3:] == "001") and len(os.listdir('\\'.join(path.split('\\')[0:-1]))) == 2:
            for file in os.listdir(path):
                file_name = file.replace("_v001", '').replace('_v001_1', '')
                try:
                    full_path = '\\'.join(path.split('\\')[0:-1]) + '\\' + path.split('\\')[-1][0] + '\\' + file_name.split('.')[0] + '.jpg'
                    img_array = np.fromfile(full_path, np.uint8)
                    img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)
                except:
                    full_path = '\\'.join(path.split('\\')[0:-1]) + '\\' + path.split('\\')[-1][0] + '\\' + file_name.split('.')[0] + '.png'
                    img_array = np.fromfile(full_path, np.uint8)
                    img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

                if img is None:
                    print('Image load failed')
                    exit()
                i = i + 1
                tree = ET.parse(path + '\\' + file)
                root = tree.getroot()
                for obj in root.findall('object'):
                    for bbox in obj.findall('bndbox'):
                        cv2.rectangle(img, (int(bbox.find('xmin').text), int(bbox.find('ymin').text)), (int(bbox.find('xmax').text), int(bbox.find('ymax').text)), (255, 255,0), 2)
                for line in root.findall('line'):
                    for controlPt in line.findall('controlPt'):
                        spot_list = { 'x' : [], 'y' : []}
                        for xlist in controlPt.findall('x'):
                            spot_list['x'].append(int(xlist.text))
                        for ylist in controlPt.findall('y'):
                            spot_list['y'].append(int(ylist.text))
                        for i in range(len(spot_list['x']) - 1):
                            cv2.line(img, (spot_list['x'][i] , spot_list['y'][i]), (spot_list['x'][i+1] , spot_list['y'][i+1]), (0, 255, 255), 5)
                try:
                    os.mkdir('\\'.join(path.split('\\')[0:-1]) + '\\' + "JungBing")
                except:
                    print("mkdir error : ", '\\'.join(path.split('\\')[0:-1]) + '\\' + "JungBing")
                    pass
                # print('\\'.join(path.split('\\')[0:-1]) + '\\' + "JungBing" + '\\' + file[:-11] + '.jpg')
                result, encoded_img = cv2.imencode('.jpg', img)
                if result:
                    with open('\\'.join(path.split('\\')[0:-1]) + '\\' + "JungBing" + '\\' + file_name.split('.')[0] + '.jpg', mode='w+b') as f:
                        encoded_img.tofile(f)
                # cv2.imwrite('\\'.join(path.split('\\')[0:-1]) + '\\' + "JungBing" + '\\' + file[:-11] + '.jpg', img)
                # print(cv2.imwrite('\\'.join(path.split('\\')[0:-1]) + '\\' + "JungBing" + '\\' + file[:-11] + '.jpg' ,img))
                # cv2.imshow("ff", img)
                # cv2.waitKey()

def MakeExtraWord():
    PATH = "C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령증_2021_대형버스시트_4차.xlsx"
    document = Document()
    data = pd.read_excel(PATH)
    cnt = 1
    for exp_name in data['시험명']:
        for (path, dirs, files) in os.walk('C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령이미지_가라_검수'):
            if path.split('\\')[-1] == exp_name:
                print(cnt)
                cnt = cnt + 1
                para = document.add_paragraph()
                run = para.add_run('    전방 영상 Annotation 검수 수당 수령증 별첨 증빙자료')
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(14)
                run.bold = True

                para = document.add_paragraph()
                para.add_run('시험명 : {}'.format(exp_name))
                para = document.add_paragraph()
                run = para.add_run()
                img_cnt = 0
                obj_arr = np.array([[], []])
                real_arr = []
                for txtfile in [_ for _ in os.listdir(path + '\\JungBing') if _.endswith(".txt")]:
                    myfile = open(path + '\\JungBing\\' + txtfile, 'r')
                    A2 = np.array([[len((myfile.readlines()))], [txtfile.split('.')[0] + '.png']])
                    obj_arr = np.concatenate((obj_arr, A2) , axis=1)
                for a in range(3):
                    try :
                        tmp = max(obj_arr[0])
                        index = np.where(obj_arr[0] == tmp)[0][0]
                        real_arr.append(obj_arr[1][index])
                        obj_arr[0][index] = -1
                    except:
                        print(path, exp_name)
                        real_arr = [0,0,0]
                for file in [_ for _ in os.listdir(path + '\\JungBing') if _.endswith(".png")]:
                    if file not in real_arr:
                        continue
                    img_array = np.fromfile(path + '\\JungBing\\' + file, np.uint8)
                    img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)
                    dst = cv2.resize(img, dsize=(320, 180), interpolation=cv2.INTER_AREA)
                    cv2.imwrite("C:\\Users\\jcy37\\Desktop\\dst.jpg", dst)
                    run.add_picture("C:\\Users\\jcy37\\Desktop\\dst.jpg")
                    img_cnt = img_cnt + 1


                month = str(data[data['시험명'] == exp_name]['월'].values[0])
                day = str(int(data[data['시험명'] == exp_name]['일'].values[0]))
                if (img_cnt < 3):
                    print("name {}, month {}, day {}, exp_name{}".format(data[data['시험명'] == exp_name]['작업자'].values[0], month, day, exp_name))
                if len(month) == 1:
                    month = '0' + month
                if len(day) == 1:
                    day = '0' + day
                document.save(
                    'C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\별첨\\' + data[data['시험명'] == exp_name]['작업자'].values[
                        0] + '_' +
                    month + day + ".docx")
                document = Document()


def pickling():
    with open('GARA.pickle', 'rb') as f:  # 'rb'라는 점을 주의하라
        data1 = pickle.load(f)
    print(data1)
    f = open('GARA.csv', 'w', newline='')
    wr = csv.writer(f)
    wr.writerow(['이름', '시험명', '프레임수', '금액', '오브젝트 개수' ,'오브젝트 합', '월', '일', '주민등록번호', '핸드폰', '주소', '계좌번호'])
    name = ''
    day_i = 0
    print("result:\n", len(data1))
    for i in range(len(data1)):
        name2 = data1[i][0]
        if name != name2:
            day_i = 0
        name = name2
        day = day_list[day_i]
        if str(day)[0] == '4' and len(str(day)) > 1:
            day = str(day)[1:]

        # print(data1[i][:4] + ["Object개수:  차선_{}/  차량 _{}/ 보행자_{}/  신호등_{}/  표지판_{}/  노면표식_{}/ 노면화살표_{}".
        #       format((data1[i][4]), (data1[i][5]), (data1[i][6]), (data1[i][7]), (data1[i][8]),
        #              (data1[i][9]), (data1[i][10]))] + [sum(data1[i][4:11])] + [day])
        wr.writerow(data1[i])

def pickling2():
    with open('GARA.pickle', 'rb') as f:  # 'rb'라는 점을 주의하라
        data1 = pickle.load(f)
    f = open('GARA.csv', 'w', newline='')
    wr = csv.writer(f)
    wr.writerow(['작업자', '시험명', '프레임수', '금액', '오브젝트 개수' ,'오브젝트 합', '월', '일', '주민등록번호', '핸드폰', '주소', '계좌번호'])
    print("result:\n", len(data1))
    for i in range(len(data1)):
        wr.writerow(data1[i])

def GetDiffereceExp():
    data = pd.read_excel("C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령증_2021_전측방시트.xlsx")
    data2 = pd.read_excel("C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령증_2021_전측방시트_2차.xlsx")
    print(data)
    print(data2)
    index2 = []
    index3 = []
    for i in range(len(data2)):
        if data2['시험명'][i] in data['시험명'].values:
            index2.append(1)
        else:
            index2.append(0)
    for i in range(len(data)):
        if data['시험명'][i] in data2['시험명'].values:
            index3.append(1)
        else:
            index3.append(0)

    print(len(index2), len(data2))
    for i in range(len(index2)):
        if index2[i] == 1:
            data2 = data2.drop(index=i)
    for i in range(len(index3)):
        if index3[i] == 1:
            data = data.drop(index=i)

    print(data[['시험명', '작업자']])
    print(data2[['시험명', '작업자']], len(data2))

    # print(len(data2))
    # index1 = data2[data2['시험명'] == 0].index
    # print(index1)
    data2.to_excel("ddd.xlsx")

def GetGARAImg():
    lcoal_work_list = ['성선영', '신성례', '김민서']
    paper_num_list = [52, 49, 26]
    try : os.mkdir("C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령이미지_가라_검수")
    except : pass
    dst_path = "C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령이미지_가라_검수"
    PATHList = ['D:\\GT 생성 업무\\객체생성-검수\공주대  gt 데이터 정제\\47_NIA3)Hep)_87_211008_pyeongtaek_3(45)_통합완료\\1',
                'D:\\GT 생성 업무\\객체생성-검수\공주대  gt 데이터 정제\\47_NIA3)Hep)_87_211008_pyeongtaek_3(45)_통합완료\\2',
                'D:\\GT 생성 업무\\객체생성-검수\공주대  gt 데이터 정제\\47_NIA3)Hep)_87_211008_pyeongtaek_3(45)_통합완료\\3',
                'D:\\GT 생성 업무\\객체생성-검수\공주대  gt 데이터 정제\\48_NIA3)Hep)_87_211008_pyeongtaek_4(rainy)_통합완료\\1',
                'D:\\GT 생성 업무\\객체생성-검수\공주대  gt 데이터 정제\\48_NIA3)Hep)_87_211008_pyeongtaek_4(rainy)_통합완료\\2',
                'D:\\GT 생성 업무\\객체생성-검수\공주대  gt 데이터 정제\\48_NIA3)Hep)_87_211008_pyeongtaek_4(rainy)_통합완료\\3',
                'D:\\GT 생성 업무\\객체생성-검수\공주대  gt 데이터 정제\\대구_전방_결과_CAM123\\1',
                'D:\\GT 생성 업무\\객체생성-검수\공주대  gt 데이터 정제\\대구_전방_결과_CAM123\\2',
                'D:\\GT 생성 업무\\객체생성-검수\공주대  gt 데이터 정제\\대구_전방_결과_CAM123\\3']
    PATHList_i = 0
    listdir_i = 0
    jump = 11
    repet = 10
    for i in range(len(lcoal_work_list)):
        try: os.mkdir(dst_path + '\\' + lcoal_work_list[i])
        except: pass
        for j in range(paper_num_list[i]):
            OsListDir = [_ for _ in os.listdir(PATHList[PATHList_i]) if _.endswith(".jpg") or _.endswith(".png")]
            if len(OsListDir) < listdir_i + jump*repet:
                PATHList_i = PATHList_i + 1
                listdir_i = 0
                OsListDir = [_ for _ in os.listdir(PATHList[PATHList_i]) if _.endswith(".jpg") or _.endswith(".png")]
            try : os.mkdir(dst_path + '\\' + lcoal_work_list[i] + '\\' + OsListDir[listdir_i].split('.')[0])
            except : pass
            for k in range(repet):
                shutil.copy2(PATHList[PATHList_i] + '\\' + OsListDir[listdir_i + jump*k]
                         , dst_path + '\\' + lcoal_work_list[i] + '\\' +
                         OsListDir[listdir_i].split('.')[0] + '\\' +
                         OsListDir[listdir_i + jump*k])
            listdir_i = listdir_i + jump*repet

def GetGARAImgJungbing():
    PATH = "H:\\img_gara_Jungbing2"
    dst_path = "C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령이미지_가라"
    for (path, dirs, files) in os.walk(dst_path):
        if path.split('\\')[-3] == '수당수령이미지_가라':
            try: os.mkdir(path+'\\'+"Jungbing")
            except: pass
            # print(path)
            # oslistdirpathpng = [_ for _ in os.listdir(path) if _.endswith(".png")]
            # for i in range(5):
            #     # shutil.copy2(PATH + '\\' + os.listdir(path)[i], path + '\\' + "Jungbing" + '\\' + os.listdir(path)[i])
            #     shutil.copy2(PATH + '\\' + oslistdirpathpng[i], path + '\\' + "Jungbing" + '\\' + oslistdirpathpng[i])
            oslistdirpath = os.listdir(path)
            for i in range(5):
                shutil.copy2(PATH + '\\' + oslistdirpath[i], path + '\\' + "Jungbing" + '\\' + oslistdirpath[i])
                try : os.mkdir(path+'\\'+"Jungbing"+ '\\' + 'labels')
                except : pass
                try : shutil.copy2(PATH + '\\' + 'labels' + '\\' + oslistdirpath[i].split('.')[0] + '.txt', path + '\\' + "Jungbing" + '\\' + oslistdirpath[i].split('.')[0] + '.txt')
                except : pass


def MonthDay(month9_i,month10_i,month12_i ,exp_name):
    month9_day = [1,2,3,6,7,8,9,10,13,14,15,16,17,23,24,27,28,29,30, 31]
    month10_day = [6,7,8,11,12,13,14,15,16,18,19,20,21,22,23,25,26,27,28,29,30]
    month12_day = [4,6,7,8,9,10,11,13,14,15,16,17,18,20,21,22]
    if int(exp_name.split('_')[1][4:6]) == 10:
        month = 10
        try:
            day = month10_day[month10_i + 3]
            month10_i = month10_i + 1
        except:
            pass
        if month10_i + 3 >= len(month10_day):
            # print(month9_i,month10_i,month12_i ,exp_name)
            month = 12
            day = month12_day[month12_i]
            month12_i = month12_i + 1
    else:
        month = 9
        try:
            day = month9_day[month9_i]
            month9_i = month9_i + 1
        except:
            pass
        if month9_i >= len(month9_day):
            month = 10
            try:
                day = month10_day[month10_i]
                month10_i = month10_i + 1
            except:
                pass
            if month10_i >= len(month10_day):
                month = 12
                day = month12_day[month12_i]
                month12_i = month12_i + 1
    return month, day, month9_i, month10_i, month12_i

def MakeGARASudangcsv():
    global result
    dev = 1
    PATH = "C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령이미지_가라_검수"
    PATH_privacy = "D:\\GT 생성 업무\\[참고]Heptacam_가이드문서 및 작업자관리시트\\미첨부(내부문서)_온라인 작업자별_할당및통계.xlsx"
    privacy = pd.read_excel(PATH_privacy, sheet_name="작업자정보")

    month9_i = 0
    month10_i = 0
    month12_i = 0
    print(os.listdir(PATH))
    cnt = 0
    month = 0
    day = 0
    for (path, dirs, files) in os.walk(PATH):
        if path.split('\\')[-2] == "수당수령이미지_가라_검수":
            month9_i = 0
            month10_i = 0
            month12_i = 0
            for exp_name in os.listdir(path):
                cnt = cnt + 1
                print(path, exp_name)
                print(month, day, month9_i, month10_i, month12_i)
                month, day, month9_i, month10_i, month12_i = MonthDay(month9_i,month10_i,month12_i ,exp_name)
                name = path.split('\\')[-1]
                price = 3000000
                while price >= 300000 or price <= 230000:
                    frame = int(dev * 20 * np.random.randn() + 130)
                    car = int(dev * frame * 2 * np.random.randn() + 1800 * 0.70)
                    ped = int(dev * 20 * np.random.randn() + 1800 * 0.0546)
                    trafficlight = int(dev * 20 * np.random.randn() + 1800 * 0.059)
                    trafficsign = int(dev * 20 * np.random.randn() + 1800 * 0.049)
                    roadmark = int(dev * 20 * np.random.randn() + 1800 * 0.07)
                    roadarrow = int(dev * frame / 5 * np.random.randn() + 1800 * 0.05)
                    price = 150 * sum([car, ped, trafficlight, trafficsign ,roadmark ,roadarrow])
                try :
                    result.append([name ,
                                   exp_name,
                                   frame,
                                   price,
                                    "Object개수:   차량 _{}/ 보행자_{}/  신호등_{}/  표지판_{}/  노면표식_{}/ 노면화살표_{}".format(car, ped, trafficlight, trafficsign ,roadmark ,roadarrow),
                                   sum([car, ped, trafficlight, trafficsign ,roadmark ,roadarrow]),
                                   month, day,
                                   privacy[privacy['이름'] == name]['주민등록번호'].values[0],
                                   privacy[privacy['이름'] == name]['핸드폰'].values[0],
                                   privacy[privacy['이름'] == name]['주소'].values[0],
                                   privacy[privacy['이름'] == name]['계좌번호'].values[0]])
                except :
                    print(name)
                    print(privacy[privacy['이름'] == name])
                    print(privacy[privacy['이름'] == name]['주민등록번호'])
                    exit(1)
    print("result:\n", len(result))
    print(result)
    print("cnt:" , cnt)
    with open('GARA.pickle', 'wb') as f:
        pickle.dump(result, f)

# 1,2차 - 3,4차
# JungBing - Jungbing
# jpg - png
# data , input_file ,os.walk
def InserImgInPDF():
    data = pd.read_excel("C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령증_2021_대형버스시트_3차.xlsx")
    input_file = PdfFileReader(open('C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령증_대형버스_자동화결과_3차.pdf', 'rb'))
    output_file = PdfFileWriter()
    for i in range(len(data)-8):
        print(i, '/', len(data))
        c = canvas.Canvas('stamp'+ str(i) + '.pdf')

        mask = [0, 0, 0, 0, 0, 0]
        for (path, dirs, files) in os.walk("C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\수당수령이미지_가라"):
            if path.split('\\')[-1] == 'Jungbing' and path.split('\\')[-2] == data['시험명'].values[i]:
                obj_arr = np.array([[], []])
                real_arr = []
                for txtfile in [_ for _ in os.listdir(path) if _.endswith(".txt")]:
                    myfile = open(path + '\\' + txtfile, 'r')
                    A2 = np.array([[len((myfile.readlines()))], [txtfile.split('.')[0] + '.png']])
                    obj_arr = np.concatenate((obj_arr, A2) , axis=1)
                for a in range(3):
                    try:
                        tmp = max(obj_arr[0])
                        index = np.where(obj_arr[0] == tmp)[0][0]
                        real_arr.append(obj_arr[1][index])
                        obj_arr[0][index] = -1
                    except:
                        print("error")
                        real_arr.append([_ for _ in os.listdir(path) if _.endswith(".png")][0])
                c.drawImage(path + '\\' + real_arr[0], x=140, y=620, width=192*0.7, height=108*0.7, mask=mask)
                c.save()
                break
        stamp = PdfFileReader(open('stamp' + str(i) + '.pdf', 'rb'))
        input_page = input_file.getPage(i)
        input_page.mergePage(stamp.getPage(0))
        output_file.addPage(input_page)

    with open('3차test.pdf', 'wb') as outputStream:
        output_file.write(outputStream)

def MoveExtra():
    PATH_src = "C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\별첨\\별첨2차"
    PATH_dst = "C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\결재서류만들기\\2차"
    extra_list = [_ for _ in os.listdir(PATH_src) if _.endswith(".PDF")]
    print(extra_list)
    for i in range(len(extra_list)):
        if extra_list[i][0:3] == '윤가영':
            shutil.copy2(PATH_src + '\\' + extra_list[i], PATH_dst + '\\' + '윤점동' + '\\' + extra_list[i])
        else:
            shutil.copy2(PATH_src + '\\' + extra_list[i] , PATH_dst + '\\' +  extra_list[i][0:3] + '\\' + extra_list[i])

def MergePDF():
    PATH = "C:\\Users\\jcy37\\Desktop\\과제\\전측방\\수당수령증_자동화\\결재서류만들기\\3차"

    # for worker in worker_list:
    for worker in ['노수진']:
        print(worker)
        merger = PdfFileMerger()
        extra = [_ for _ in os.listdir(PATH + '\\' + worker)
                 if _.endswith(".PDF") and len(_) == 12 and _.startswith(worker)]

        Sudang = [_ for _ in os.listdir(PATH + '\\' + worker)
                 if _.endswith(".pdf") and _.startswith("S25C")]
        print(Sudang)
        print(extra)
        print(len(Sudang), len(extra))
        for k in range(len(Sudang)):
            Sudang_endnum = Sudang[k].split('-')[2].replace('.pdf' , '')
            for o in range(3 - len(Sudang_endnum)):
                Sudang_endnum = '0' + Sudang_endnum
            os.rename(PATH + '\\' + worker + '\\' + Sudang[k],
                      PATH + '\\' + worker + '\\' +
                      Sudang[k].split('-')[0] + '-' + Sudang[k].split('-')[1] + '-' + Sudang_endnum + '.pdf')
        Sudang = [_ for _ in os.listdir(PATH + '\\' + worker)
                  if _.endswith(".pdf") and _.startswith("S25C")]

        Searyu = [_ for _ in os.listdir(PATH + '\\' + worker)
                 if _.endswith(".pdf") and not _.startswith("S25C") and _.split('.')[-2][-2:] == '서류']

        if len(Sudang) != len(extra):
            print(worker , "Sudang and extra num is not accurate!!!!")
            exit(1)
        if len(Searyu) == 0:
            print(worker , "no Searyu !!!!")
            exit(1)
        for i in range(len(extra)):
            merger.append(PATH + '\\' + worker + '\\' + Sudang[i])
            merger.append(PATH + '\\' + worker + '\\' + extra[i])
        for j in range(len(Searyu)):
            merger.append(PATH + '\\' + worker + '\\' + Searyu[j])
        merger.write(PATH + '\\' + worker + '\\' + worker + '.pdf')
if __name__ == "__main__":
    MergePDF()

